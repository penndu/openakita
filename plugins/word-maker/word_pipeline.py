"""Deterministic document generation pipeline for word-maker."""

from __future__ import annotations

import json
import time
from collections.abc import Awaitable, Callable
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from word_brain_helper import WordBrainHelper
from word_source_loader import load_source
from word_task_manager import WordTaskManager
from word_template_engine import extract_template_vars, render_template

Emit = Callable[[str, dict[str, Any]], Awaitable[None] | None]


@dataclass(slots=True)
class WordPipelineContext:
    project_id: str
    task_dir: Path
    requirement: str = ""
    doc_type: str = "research_report"
    template_path: Path | None = None
    source_paths: list[Path] = field(default_factory=list)
    fields: dict[str, Any] = field(default_factory=dict)
    outline: dict[str, Any] = field(default_factory=dict)
    doc_markdown: str = ""
    output_path: Path | None = None
    audit: dict[str, Any] = field(default_factory=dict)
    error_kind: str | None = None
    error_message: str | None = None

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        for key in ("task_dir", "template_path", "output_path"):
            if data[key] is not None:
                data[key] = str(data[key])
        data["source_paths"] = [str(path) for path in self.source_paths]
        return data


async def _emit(emit: Emit | None, event: str, payload: dict[str, Any]) -> None:
    if emit is None:
        return
    result = emit(event, payload)
    if result is not None:
        await result


def _outline_to_markdown(outline: dict[str, Any], fields: dict[str, Any]) -> str:
    title = outline.get("title") or fields.get("title") or "文档初稿"
    chunks = [f"# {title}"]
    for section in outline.get("sections") or []:
        chunks.append(f"\n## {section.get('title', '未命名章节')}")
        goal = section.get("goal")
        if goal:
            chunks.append(str(goal))
        for bullet in section.get("bullets") or []:
            chunks.append(f"- {bullet}")
    return "\n".join(chunks).strip() + "\n"


def audit_output(path: Path | None, *, markdown: str = "", missing: list[str] | None = None) -> dict[str, Any]:
    warnings: list[str] = []
    errors: list[str] = []
    if path is None or not path.exists():
        errors.append("Output file was not created")
    elif path.stat().st_size < 100:
        errors.append("Output file is unexpectedly small")
    if missing:
        errors.append(f"Missing template variables: {', '.join(missing)}")
    if markdown and "##" not in markdown:
        warnings.append("Document has no second-level headings")
    return {"ok": not errors, "warnings": warnings, "errors": errors}


def build_ppt_asset_metadata(
    *,
    project: dict[str, Any],
    outline: dict[str, Any] | None = None,
    doc_markdown: str = "",
    export_docx_path: str | None = None,
) -> dict[str, Any]:
    """Build the stable payload consumed by future ppt-maker handoff code."""

    return {
        "project_id": project.get("id"),
        "doc_type": project.get("doc_type"),
        "title": project.get("title"),
        "outline_json": outline or {},
        "summary_md": doc_markdown[:12000],
        "export_docx_path": export_docx_path or project.get("output_path"),
        "source": "word-maker",
        "handoff_version": 1,
    }


async def run_pipeline(
    ctx: WordPipelineContext,
    *,
    manager: WordTaskManager,
    brain_helper: WordBrainHelper | None = None,
    emit: Emit | None = None,
) -> WordPipelineContext:
    started = time.time()
    try:
        await manager.update_project_safe(ctx.project_id, status="rendering")
        await _emit(emit, "project_update", {"project_id": ctx.project_id, "status": "rendering"})

        source_texts: list[str] = []
        for source_path in ctx.source_paths:
            result = load_source(source_path)
            if result.ok:
                source_texts.append(result.text)
                await manager.add_source(
                    ctx.project_id,
                    source_type=result.source_type,
                    filename=source_path.name,
                    path=str(source_path),
                    text_preview=result.text[:1200],
                    parse_status="parsed",
                )
            else:
                await manager.add_source(
                    ctx.project_id,
                    source_type=result.source_type,
                    filename=source_path.name,
                    path=str(source_path),
                    parse_status="failed",
                    error_message=result.error,
                )

        sources_text = "\n\n".join(source_texts)
        if not ctx.outline and brain_helper is not None:
            outline_result = await brain_helper.generate_outline(
                requirement=ctx.requirement,
                doc_type=ctx.doc_type,
                sources_text=sources_text,
            )
            ctx.outline = outline_result.data
        if not ctx.outline:
            ctx.outline = {
                "title": ctx.fields.get("title") or "文档初稿",
                "sections": [
                    {"id": "main", "title": "正文", "goal": ctx.requirement or "整理核心内容", "bullets": []}
                ],
                "missing_inputs": [],
            }

        ctx.doc_markdown = _outline_to_markdown(ctx.outline, ctx.fields)
        drafts_dir = ctx.task_dir / "drafts"
        exports_dir = ctx.task_dir / "exports"
        drafts_dir.mkdir(parents=True, exist_ok=True)
        exports_dir.mkdir(parents=True, exist_ok=True)
        (drafts_dir / "draft.md").write_text(ctx.doc_markdown, encoding="utf-8")

        missing: list[str] = []
        if ctx.template_path is not None:
            inspection = extract_template_vars(ctx.template_path, context=ctx.fields)
            missing = inspection.missing
            ctx.output_path = exports_dir / "document.docx"
            render = render_template(ctx.template_path, ctx.output_path, ctx.fields)
            if not render.ok:
                raise RuntimeError(render.error or "Template render failed")
            await manager.add_template(
                ctx.project_id,
                label=ctx.template_path.name,
                path=str(ctx.template_path),
                variables=inspection.variables,
                validation={"missing": inspection.missing, "engine": inspection.engine},
            )
        else:
            from docx import Document

            document = Document()
            for line in ctx.doc_markdown.splitlines():
                if line.startswith("# "):
                    document.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    document.add_heading(line[3:], level=2)
                elif line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                elif line.strip():
                    document.add_paragraph(line)
            ctx.output_path = exports_dir / "document.docx"
            document.save(ctx.output_path)

        ctx.audit = audit_output(ctx.output_path, markdown=ctx.doc_markdown, missing=missing)
        (ctx.task_dir / "audit.json").write_text(
            json.dumps(ctx.audit, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        version = await manager.add_draft_version(
            ctx.project_id,
            outline=ctx.outline,
            fields=ctx.fields,
            doc_markdown=ctx.doc_markdown,
            export_path=str(ctx.output_path),
            audit=ctx.audit,
        )
        status = "succeeded" if ctx.audit.get("ok") else "failed"
        await manager.update_project_safe(
            ctx.project_id,
            status=status,
            output_path=str(ctx.output_path) if ctx.output_path else None,
            completed_at=time.time(),
            metadata={"pipeline_ms": int((time.time() - started) * 1000), "version": version["version"]},
        )
        await _emit(emit, "project_update", {"project_id": ctx.project_id, "status": status})
    except Exception as exc:
        ctx.error_kind = "pipeline"
        ctx.error_message = str(exc)
        await manager.update_project_safe(
            ctx.project_id,
            status="failed",
            error_kind=ctx.error_kind,
            error_message=ctx.error_message,
            completed_at=time.time(),
        )
        await _emit(
            emit,
            "project_update",
            {"project_id": ctx.project_id, "status": "failed", "error": ctx.error_message},
        )
    finally:
        (ctx.task_dir / "metadata.json").write_text(
            json.dumps(ctx.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    return ctx

