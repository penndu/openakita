from __future__ import annotations

from pathlib import Path

from docx import Document
from word_template_engine import extract_template_vars, render_template


def _make_template(path: Path) -> None:
    document = Document()
    document.add_heading("{{ title }}", level=1)
    document.add_paragraph("客户：{{ company }}")
    document.add_paragraph("摘要：{{ summary }}")
    document.save(path)


def test_extract_template_vars_with_missing_context(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _make_template(template)

    inspection = extract_template_vars(template, context={"title": "报告"})

    assert inspection.engine in {"docxtpl", "regex"}
    assert set(inspection.variables) >= {"title", "company", "summary"}
    assert inspection.missing == ["company", "summary"]
    assert inspection.ok is False


def test_render_template_rejects_missing_vars(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    output = tmp_path / "out.docx"
    _make_template(template)

    result = render_template(template, output, {"title": "报告"})

    assert result.ok is False
    assert result.missing == ["company", "summary"]
    assert not output.exists()


def test_render_template_happy_path(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    output = tmp_path / "out.docx"
    _make_template(template)

    result = render_template(
        template,
        output,
        {
            "title": "验收报告",
            "company": "OpenAkita",
            "summary": "项目已完成验收。",
        },
    )

    assert result.ok is True
    assert output.exists()
    rendered = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "验收报告" in rendered
    assert "OpenAkita" in rendered
    assert "{{" not in rendered

