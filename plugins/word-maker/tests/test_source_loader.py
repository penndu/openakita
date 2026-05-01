from __future__ import annotations

from pathlib import Path

from docx import Document
from word_source_loader import load_source


def test_load_text_source(tmp_path: Path) -> None:
    source = tmp_path / "notes.md"
    source.write_text("# 标题\n正文内容", encoding="utf-8")

    result = load_source(source)

    assert result.ok is True
    assert result.source_type == "md"
    assert "正文内容" in result.text


def test_load_csv_source(tmp_path: Path) -> None:
    source = tmp_path / "data.csv"
    source.write_text("name,value\nA,1\n", encoding="utf-8")

    result = load_source(source)

    assert result.ok is True
    assert result.source_type == "csv"
    assert "name | value" in result.text


def test_load_docx_source(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("项目背景")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "结果"
    document.save(source)

    result = load_source(source)

    assert result.ok is True
    assert result.source_type == "docx"
    assert "项目背景" in result.text
    assert "指标 | 结果" in result.text


def test_unsupported_source_returns_error(tmp_path: Path) -> None:
    source = tmp_path / "archive.bin"
    source.write_bytes(b"binary")

    result = load_source(source)

    assert result.ok is False
    assert "Unsupported" in result.error

