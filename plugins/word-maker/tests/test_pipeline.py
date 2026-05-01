from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from word_pipeline import WordPipelineContext, audit_output, build_ppt_asset_metadata, run_pipeline
from word_task_manager import WordTaskManager


@pytest.mark.asyncio
async def test_pipeline_generates_docx_without_template(tmp_path: Path) -> None:
    manager = WordTaskManager(tmp_path / "word-maker.db", tmp_path / "projects")
    async with manager:
        project = await manager.create_project({"title": "调研报告", "doc_type": "research_report"})
        ctx = WordPipelineContext(
            project_id=project["id"],
            task_dir=manager.project_dir(project["id"]),
            requirement="生成调研报告",
            outline={
                "title": "调研报告",
                "sections": [
                    {"id": "background", "title": "背景", "goal": "说明背景", "bullets": ["目标明确"]}
                ],
            },
        )

        result = await run_pipeline(ctx, manager=manager)

        assert result.output_path is not None
        assert result.output_path.exists()
        assert result.audit["ok"] is True
        assert (await manager.get_project(project["id"]))["status"] == "succeeded"
        assert (manager.project_dir(project["id"]) / "metadata.json").exists()


@pytest.mark.asyncio
async def test_pipeline_uses_template_and_records_failure(tmp_path: Path) -> None:
    manager = WordTaskManager(tmp_path / "word-maker.db", tmp_path / "projects")
    async with manager:
        project = await manager.create_project({"title": "模板报告", "doc_type": "research_report"})
        template = tmp_path / "template.docx"
        document = Document()
        document.add_paragraph("标题：{{ title }}")
        document.add_paragraph("客户：{{ company }}")
        document.save(template)
        ctx = WordPipelineContext(
            project_id=project["id"],
            task_dir=manager.project_dir(project["id"]),
            template_path=template,
            fields={"title": "报告"},
        )

        result = await run_pipeline(ctx, manager=manager)

        assert result.error_kind == "pipeline"
        assert "Missing template variables" in (result.error_message or "")
        assert (await manager.get_project(project["id"]))["status"] == "failed"


def test_audit_output_reports_missing_file() -> None:
    audit = audit_output(Path("missing.docx"), markdown="# Title")

    assert audit["ok"] is False
    assert audit["errors"]


def test_build_ppt_asset_metadata() -> None:
    metadata = build_ppt_asset_metadata(
        project={"id": "doc_1", "doc_type": "research_report", "title": "调研报告"},
        outline={"title": "调研报告"},
        doc_markdown="# 调研报告",
        export_docx_path="out.docx",
    )

    assert metadata["project_id"] == "doc_1"
    assert metadata["handoff_version"] == 1
    assert metadata["export_docx_path"] == "out.docx"

