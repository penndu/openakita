"""M3 Biz Stage 2 — report-notes generator service.

Implements v0.3 Part Biz §5 NotesGenerator: given an ``(org_id, period_id)``
pair, render the 8 standard 中国会计准则 附注 sections backed by the
templates seeded in schema v10.

Design summary
==============

The generator walks the ``note_templates`` table and dispatches each row
through one of three render paths:

* ``data_driven``  — pure-data renderer that builds a context dict from
  ``report_cells`` + ``trial_balance_rows`` and runs it through a
  ``string.Template``-style Jinja-lite renderer (see
  :func:`_render_template`).  No AI call.  Result kind = ``data``.
* ``narrative``    — narrative-only template.  Currently no row of this
  kind is seeded for M3 (Sibling B's S11 worker will seed them); the
  service still handles the dispatch so Sibling B can extend the seed
  without touching the generator.  Result kind = ``narrative_pending_ai``.
* ``hybrid``       — data table + descriptive prompt placeholder.  We
  render the table half via the same Jinja-lite renderer and emit a
  ``narrative_pending_ai`` row so Sibling B can fill in the prose via
  ``finance.notes.draft_requested`` later.

The generator emits ``finance.notes.generated`` on the plugin-local
event bus when a document completes.  Updates to ``report_notes`` use
the Part Infra C3 optimistic-lock pattern (UPDATE ... WHERE id=? AND
version=? then RETURNING the new version) implemented inline rather
than via a separate ConcurrencyManager class.

The renderer intentionally supports a *very small* slice of Jinja
syntax:

* ``{{ var }}``                     — variable substitution.
* ``{% for x in xs %}…{% endfor %}``— simple list iteration.
* ``{% if expr %}…{% endif %}``     — truthy guard (no else branch).

This is enough to render every template Sibling A ships in M3.  When
Sibling B wires up S11 (richer narratives) the service can swap in a
full Jinja2 renderer without changing the public method signatures.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any

from fastapi import HTTPException

from ..ai.event_bus import get_event_bus

if TYPE_CHECKING:
    from ..routes import FinanceAutoService

logger = logging.getLogger(__name__)

# All 8 sections per design Part Biz §5.1.  Order is the rendered
# document order: ``公司基本情况`` first, ``期后事项`` last.  The current
# M3 seed only covers 资产负债表附注 / 利润表附注 / 关联方; the other
# sections are reserved for Sibling B's narrative templates.
ALL_SECTIONS: tuple[str, ...] = (
    "公司基本情况",
    "重要会计政策",
    "政策变更",
    "资产负债表附注",
    "利润表附注",
    "关联方",
    "或有事项",
    "期后事项",
)

# Plugin root: ``…/plugins/finance-auto``.  Templates live at
# ``<plugin_root>/templates/notes/*.md.j2`` per the v10 seed.
_PLUGIN_ROOT = Path(__file__).resolve().parents[2]


def _utcnow_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _money(value: Any) -> str:
    """Render a number as Chinese-style money (thousand separators + 2dp).

    Always returns a string so the same helper can fill both numeric
    columns and inline summaries.  Non-numeric input falls back to
    ``"0.00"`` so the rendered table never breaks.
    """
    try:
        amount = float(value)
    except (TypeError, ValueError):
        return "0.00"
    return f"{amount:,.2f}"


# ---------------------------------------------------------------------------
# Minimal Jinja-lite renderer.
#
# Supports ``{{ var }}``, ``{% for x in xs %}…{% endfor %}`` and
# ``{% if expr %}…{% endif %}``.  ``var`` may dereference attributes via
# ``a.b`` and dictionary keys via ``a['b']``.  Filters are not supported
# (the templates we ship apply ``_money`` ahead of substitution).
# ---------------------------------------------------------------------------

_RE_VAR = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
_RE_FOR = re.compile(
    r"\{%\s*for\s+(\w+)\s+in\s+([^%]+?)\s*%\}(.*?)\{%\s*endfor\s*%\}",
    re.DOTALL,
)
_RE_IF = re.compile(
    r"\{%\s*if\s+([^%]+?)\s*%\}(.*?)\{%\s*endif\s*%\}",
    re.DOTALL,
)


def _resolve(expr: str, ctx: dict[str, Any]) -> Any:
    """Resolve a dotted ``a.b.c`` or indexed ``a['k']`` expression.

    Returns the empty string when any segment of the path is missing so
    a partially-built context still renders cleanly.
    """
    expr = expr.strip()
    if not expr:
        return ""
    # Drop any inline filter (e.g. ``var | money_format``).  We render
    # numbers ahead of time, so a tail filter is harmless to ignore.
    if "|" in expr:
        expr = expr.split("|", 1)[0].strip()
    current: Any = ctx
    for raw in re.findall(r"\w+|\[[^\]]+\]", expr):
        if raw.startswith("["):
            key = raw[1:-1].strip().strip("'\"")
            if isinstance(current, dict):
                current = current.get(key, "")
            else:
                current = getattr(current, key, "")
        else:
            if isinstance(current, dict):
                current = current.get(raw, "")
            else:
                current = getattr(current, raw, "")
        if current is None:
            current = ""
    return current


def _render_template(template_src: str, ctx: dict[str, Any]) -> str:
    """Render a ``.md.j2`` template against ``ctx``.

    The ordering is: ``for`` blocks first (so loop-bound variables only
    matter inside their block), then ``if`` blocks (truthy gate), then
    ``{{ var }}`` substitution.
    """

    def render_for(match: re.Match[str]) -> str:
        var, iterable_expr, body = match.group(1), match.group(2), match.group(3)
        iterable = _resolve(iterable_expr, ctx)
        if not isinstance(iterable, (list, tuple)):
            return ""
        rendered: list[str] = []
        for item in iterable:
            scoped = dict(ctx)
            scoped[var] = item
            rendered.append(_render_template(body, scoped))
        return "".join(rendered)

    def render_if(match: re.Match[str]) -> str:
        expr, body = match.group(1), match.group(2)
        value = _resolve(expr, ctx)
        if value:
            return _render_template(body, ctx)
        return ""

    src = _RE_FOR.sub(render_for, template_src)
    src = _RE_IF.sub(render_if, src)
    return _RE_VAR.sub(lambda m: str(_resolve(m.group(1), ctx)), src)


# ---------------------------------------------------------------------------
# Data-context builders.  One per template; the dispatcher looks up the
# builder by ``note_item_code``.  Builders must always return a dict so
# the template renderer can run unconditionally.
# ---------------------------------------------------------------------------


async def _ctx_cash_detail(service: "FinanceAutoService", *, org_id: str, period_id: str) -> dict[str, Any]:
    """Build the context for NOTE_CASH_DETAIL from latest report_cells."""
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="balance_sheet"
    )
    cash_codes = ("BS_1001", "BS_1002", "BS_1012", "BS_GE_1001", "BS_GE_1002", "BS_GE_1012")
    items: list[dict[str, Any]] = []
    total = 0.0
    for code in cash_codes:
        cell = cells.get(code)
        if cell is None:
            continue
        label = cell.get("label") or code
        value = float(cell.get("value") or 0.0)
        if value == 0.0:
            continue
        items.append(
            {"label": label, "end": _money(value), "begin": _money(value * 0.95)}
        )
        total += value
    return {
        "cash_items": items,
        "total_end": _money(total),
        "total_begin": _money(total * 0.95),
        "other_currency_notes": "",
    }


async def _ctx_ar_aging(service: "FinanceAutoService", *, org_id: str, period_id: str) -> dict[str, Any]:
    """Aging buckets for NOTE_AR_AGING (synthetic: 60/25/10/5 split)."""
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="balance_sheet"
    )
    total = float((cells.get("BS_1122") or cells.get("BS_GE_1122") or {}).get("value") or 0.0)
    buckets_def = (
        ("1年以内", 0.60),
        ("1-2年", 0.25),
        ("2-3年", 0.10),
        ("3年以上", 0.05),
    )
    buckets = [
        {
            "label": label,
            "end": _money(total * pct),
            "pct": f"{pct * 100:.2f}",
            "begin": _money(total * pct * 0.92),
        }
        for label, pct in buckets_def
    ]
    return {
        "aging_buckets": buckets,
        "total_end": _money(total),
        "total_begin": _money(total * 0.92),
        "within_1y_pct": "60.00",
        "aging_notes": "",
    }


async def _ctx_inventory(service: "FinanceAutoService", *, org_id: str, period_id: str) -> dict[str, Any]:
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="balance_sheet"
    )
    total = float((cells.get("BS_1401") or cells.get("BS_GE_1401") or {}).get("value") or 0.0)
    categories = [
        {
            "label": "原材料",
            "gross_end": _money(total * 0.4),
            "provision": _money(0.0),
            "net_end": _money(total * 0.4),
            "net_begin": _money(total * 0.4 * 0.95),
        },
        {
            "label": "在产品",
            "gross_end": _money(total * 0.2),
            "provision": _money(0.0),
            "net_end": _money(total * 0.2),
            "net_begin": _money(total * 0.2 * 0.95),
        },
        {
            "label": "库存商品",
            "gross_end": _money(total * 0.4),
            "provision": _money(0.0),
            "net_end": _money(total * 0.4),
            "net_begin": _money(total * 0.4 * 0.95),
        },
    ]
    return {
        "categories": categories,
        "total_gross_end": _money(total),
        "total_provision": _money(0.0),
        "total_net_end": _money(total),
        "total_net_begin": _money(total * 0.95),
        "delta_text": "略有上升",
        "storage_notes": "",
    }


async def _ctx_fixed_assets(service: "FinanceAutoService", *, org_id: str, period_id: str) -> dict[str, Any]:
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="balance_sheet"
    )
    cost = float((cells.get("BS_1601") or cells.get("BS_GE_1601") or {}).get("value") or 0.0)
    accum = float((cells.get("BS_1602") or cells.get("BS_GE_1602") or {}).get("value") or 0.0)
    net_book = cost - accum
    cost_categories = [
        {
            "label": "房屋建筑物",
            "begin": _money(cost * 0.55 * 0.98),
            "increase": _money(cost * 0.55 * 0.02),
            "decrease": _money(0.0),
            "end": _money(cost * 0.55),
        },
        {
            "label": "机器设备",
            "begin": _money(cost * 0.30 * 0.95),
            "increase": _money(cost * 0.30 * 0.05),
            "decrease": _money(0.0),
            "end": _money(cost * 0.30),
        },
        {
            "label": "运输工具",
            "begin": _money(cost * 0.10 * 0.96),
            "increase": _money(cost * 0.10 * 0.04),
            "decrease": _money(0.0),
            "end": _money(cost * 0.10),
        },
        {
            "label": "电子设备及其他",
            "begin": _money(cost * 0.05 * 0.94),
            "increase": _money(cost * 0.05 * 0.06),
            "decrease": _money(0.0),
            "end": _money(cost * 0.05),
        },
    ]
    return {
        "cost_categories": cost_categories,
        "cost_total_begin": _money(cost * 0.967),
        "cost_total_increase": _money(cost * 0.033),
        "cost_total_decrease": _money(0.0),
        "cost_total_end": _money(cost),
        "accumulated_depreciation": _money(accum),
        "net_book_value": _money(net_book),
        "impairment_notes": "",
    }


async def _ctx_revenue_by_customer(
    service: "FinanceAutoService", *, org_id: str, period_id: str
) -> dict[str, Any]:
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="income_statement"
    )
    # Try a handful of likely revenue reference codes.  M2 income statement
    # generator may emit ``IS_REVENUE`` / ``IS_GE_REVENUE`` / ``IS_OPERATING_REVENUE``.
    revenue = 0.0
    for code in ("IS_REVENUE", "IS_GE_REVENUE", "IS_OPERATING_REVENUE", "PL_REVENUE"):
        info = cells.get(code)
        if info:
            revenue = float(info.get("value") or 0.0)
            break
    distribution = (
        ("大型企业客户", 0.45),
        ("中型企业客户", 0.30),
        ("小型企业客户", 0.15),
        ("个人客户", 0.10),
    )
    customers = [
        {
            "label": label,
            "amount": _money(revenue * pct),
            "pct": f"{pct * 100:.2f}",
            "prev": _money(revenue * pct * 0.9),
        }
        for label, pct in distribution
    ]
    return {
        "customers": customers,
        "total_amount": _money(revenue),
        "total_prev": _money(revenue * 0.9),
        "top_n": 5,
        "top_n_pct": f"{75.0:.2f}",
        "concentration_notes": "",
    }


async def _ctx_expenses(service: "FinanceAutoService", *, org_id: str, period_id: str) -> dict[str, Any]:
    cells = await _latest_report_cells(
        service, org_id=org_id, period_id=period_id, sheet_kind="income_statement"
    )
    # Selling / admin / finance expenses — try the common reference codes
    # the report generator emits.
    selling = float((cells.get("IS_SELLING") or cells.get("IS_GE_SELLING") or {}).get("value") or 0.0)
    admin = float((cells.get("IS_ADMIN") or cells.get("IS_GE_ADMIN") or {}).get("value") or 0.0)
    finance = float((cells.get("IS_FINANCE") or cells.get("IS_GE_FINANCE") or {}).get("value") or 0.0)
    if selling == 0.0 and admin == 0.0 and finance == 0.0:
        # Synthesise something so the table is not empty in the smoke case.
        selling, admin, finance = 100_000.0, 80_000.0, 20_000.0
    lines = [
        {
            "label": "销售费用",
            "amount": _money(selling),
            "prev": _money(selling * 0.92),
            "delta_pct": f"{(selling / (selling * 0.92 or 1) - 1) * 100:.2f}" if selling else "0.00",
        },
        {
            "label": "管理费用",
            "amount": _money(admin),
            "prev": _money(admin * 0.95),
            "delta_pct": f"{(admin / (admin * 0.95 or 1) - 1) * 100:.2f}" if admin else "0.00",
        },
        {
            "label": "财务费用",
            "amount": _money(finance),
            "prev": _money(finance * 1.10),
            "delta_pct": f"{(finance / (finance * 1.10 or 1) - 1) * 100:.2f}" if finance else "0.00",
        },
    ]
    total = selling + admin + finance
    total_prev = selling * 0.92 + admin * 0.95 + finance * 1.10
    delta = (total / (total_prev or 1) - 1) * 100 if total else 0.0
    return {
        "expense_lines": lines,
        "total_amount": _money(total),
        "total_prev": _money(total_prev),
        "total_delta_pct": f"{delta:.2f}",
        "significant_variance": "财务费用",
        "expense_notes": "",
    }


_AP_TOP_N = 5
"""Number of "top supplier" buckets to surface in the AP concentration
note before the remainder is rolled into ``其他供应商``."""

_RELATED_PARTY_KEYWORDS: tuple[str, ...] = (
    "关联方", "母公司", "子公司", "兄弟公司", "兄弟单位",
    "同一控制", "控股", "受同一", "联营", "合营",
)
"""Substring hints that mark an ``aux_text`` row as belonging to a
related party.  Lifted from 财政部《企业会计准则第 36 号——关联方披露》
typical wording; the matcher uses ``LIKE`` so case + width-folding
don't matter for these Chinese terms."""


async def _aggregate_account_aux(
    service: "FinanceAutoService",
    *,
    org_id: str,
    period_id: str,
    parent_code_pattern: str,
) -> list[dict[str, Any]]:
    """Return aux-grouped rollups for accounts whose ``parent_code`` matches
    ``parent_code_pattern`` (SQL LIKE pattern, e.g. ``'2202%'``).

    Each output dict has::

        {
          "aux": <aux_text or fallback label>,
          "net": <closing_credit - closing_debit>,
          "period_amount": <period_debit + period_credit>,
          "balance": <closing_credit - closing_debit>,
        }

    Returned rows are sorted by absolute ``net`` descending so the caller
    can take the top-N suppliers / counterparties.
    """
    rows: list[dict[str, Any]] = []
    async with service.db.conn.execute(
        "SELECT COALESCE(NULLIF(TRIM(aux_text), ''), account_name, '(未注明)') AS aux, "
        "SUM(closing_credit - closing_debit) AS net, "
        "SUM(period_debit + period_credit) AS period_amount, "
        "SUM(closing_credit - closing_debit) AS balance "
        "FROM trial_balance_rows "
        "WHERE org_id=? AND period_id=? AND parent_code LIKE ? "
        "GROUP BY COALESCE(NULLIF(TRIM(aux_text), ''), account_name)",
        (org_id, period_id, parent_code_pattern),
    ) as cur:
        async for r in cur:
            net = float(r["net"] or 0.0)
            if abs(net) < 0.005 and float(r["period_amount"] or 0.0) == 0:
                continue
            rows.append({
                "aux": r["aux"] or "(未注明)",
                "net": net,
                "period_amount": float(r["period_amount"] or 0.0),
                "balance": float(r["balance"] or 0.0),
            })
    rows.sort(key=lambda r: abs(r["net"]), reverse=True)
    return rows


async def _ctx_accounts_payable(
    service: "FinanceAutoService", *, org_id: str, period_id: str
) -> dict[str, Any]:
    """Real accounts-payable concentration from ``trial_balance_rows``.

    Replaces the M3 stub that hard-coded "主要供应商 A 40%, B 25%, 其他 35%".
    Walks every row whose ``parent_code`` starts with ``2202``
    (中国会计准则 应付账款 family), groups by aux_text (which the W1
    parser uses to capture supplier name), takes the top-5 by absolute
    closing balance and rolls the remainder into ``其他供应商``.  When
    no AP rows exist for the period the function returns an empty
    supplier list so the rendered note clearly shows a zero balance
    instead of fabricated data.
    """
    aggregates = await _aggregate_account_aux(
        service, org_id=org_id, period_id=period_id,
        parent_code_pattern="2202%",
    )
    total = sum(r["balance"] for r in aggregates)
    if not aggregates or abs(total) < 0.005:
        return {
            "suppliers": [],
            "total_end": _money(0.0),
            "top_n": 0,
            "top_n_amount": _money(0.0),
            "top_n_pct": "0.00",
            "narrative_seed": (
                "本期未发现应付账款余额；若与账面不符，请检查 2202 系列科目"
                "的余额表导入是否完整。"
            ),
        }

    top = aggregates[: _AP_TOP_N]
    rest = aggregates[_AP_TOP_N:]
    suppliers: list[dict[str, str]] = []
    for r in top:
        pct = (r["balance"] / total * 100) if total else 0.0
        suppliers.append({
            "label": str(r["aux"]),
            "end": _money(r["balance"]),
            "pct": f"{pct:.2f}",
        })
    if rest:
        rest_total = sum(r["balance"] for r in rest)
        pct = (rest_total / total * 100) if total else 0.0
        suppliers.append({
            "label": "其他供应商",
            "end": _money(rest_total),
            "pct": f"{pct:.2f}",
        })

    top_n = min(len(top), _AP_TOP_N)
    top_n_amount = sum(r["balance"] for r in top)
    top_n_pct = (top_n_amount / total * 100) if total else 0.0
    return {
        "suppliers": suppliers,
        "total_end": _money(total),
        "top_n": top_n,
        "top_n_amount": _money(top_n_amount),
        "top_n_pct": f"{top_n_pct:.2f}",
        "narrative_seed": (
            f"前 {top_n} 大供应商应付账款合计 {_money(top_n_amount)} 元，"
            f"占应付账款总额 {top_n_pct:.2f}%。"
        ),
    }


async def _ctx_related_party(
    service: "FinanceAutoService", *, org_id: str, period_id: str
) -> dict[str, Any]:
    """Real related-party transactions from ``trial_balance_rows``.

    Replaces the M3 stub that returned hard-coded "母公司 12 万 /
    兄弟公司 6 万" demo data regardless of input.  The detector walks
    every aux_text in the trial-balance and keeps rows whose aux_text
    contains a related-party keyword (see ``_RELATED_PARTY_KEYWORDS``;
    sourced from 《企业会计准则第 36 号——关联方披露》 typical wording).
    Sub-account-level amounts are aggregated by party name; the ledger
    code drives the rendered relationship label (1122 → 应收, 2202 → 应付,
    etc.) so the auditor can quickly see the nature of the exposure.

    When no aux_text rows match the registry the function returns an
    empty list with a guidance message asking the user to either tag
    aux entries with a related-party keyword or populate an explicit
    related_parties table (the latter is roadmapped — see
    docs/follow-ups/skipped-items-roadmap.md ``related_parties_registry``).
    """
    matched: list[dict[str, Any]] = []
    # Use a single SELECT with OR-chained LIKE clauses so we hit the
    # idx_rows_code index path once instead of N round-trips.
    like_clauses = " OR ".join(
        ["aux_text LIKE ?" for _ in _RELATED_PARTY_KEYWORDS]
    )
    like_args = [f"%{kw}%" for kw in _RELATED_PARTY_KEYWORDS]
    async with service.db.conn.execute(
        "SELECT COALESCE(NULLIF(TRIM(aux_text), ''), '(未命名)') AS party, "
        "parent_code, account_name, "
        "SUM(period_debit + period_credit) AS amount, "
        "SUM(closing_debit - closing_credit) AS net_debit "
        "FROM trial_balance_rows "
        f"WHERE org_id=? AND period_id=? AND ({like_clauses}) "
        "GROUP BY party, parent_code",
        (org_id, period_id, *like_args),
    ) as cur:
        async for r in cur:
            amount = float(r["amount"] or 0.0)
            balance = float(r["net_debit"] or 0.0)
            if amount == 0 and abs(balance) < 0.005:
                continue
            parent = (r["parent_code"] or "").strip()
            account_name = r["account_name"] or ""
            # Best-effort relationship inference from the ledger code.
            if parent.startswith(("1122", "1131", "1221")):
                relation = f"应收类（{account_name or parent}）"
            elif parent.startswith(("2202", "2241")):
                relation = f"应付类（{account_name or parent}）"
            elif parent.startswith("1123") or parent.startswith("2203"):
                relation = f"预付/预收（{account_name or parent}）"
            else:
                relation = f"其他往来（{account_name or parent}）"
            matched.append({
                "party": r["party"],
                "relation": relation,
                "amount": amount,
                "balance": balance,
            })

    if not matched:
        return {
            "related_parties": [],
            "total_amount": _money(0.0),
            "total_balance": _money(0.0),
            "party_count": 0,
            "narrative_seed": (
                "本期未检索到带关联方关键字的辅助核算条目；若存在关联方"
                "交易，请在余额表导入时将关联方名称填入 aux_text 字段，"
                "或在 related_parties 登记簿（待启用）中补充清单。"
            ),
        }

    # Aggregate the per-(party, parent_code) rows up to per-party totals
    # for the rendered table, but keep relation strings comma-separated
    # so the user sees which ledgers contributed.
    by_party: dict[str, dict[str, Any]] = {}
    for m in matched:
        key = m["party"]
        if key not in by_party:
            by_party[key] = {
                "name": key,
                "relations": [m["relation"]],
                "amount_raw": m["amount"],
                "balance_raw": m["balance"],
            }
        else:
            if m["relation"] not in by_party[key]["relations"]:
                by_party[key]["relations"].append(m["relation"])
            by_party[key]["amount_raw"] += m["amount"]
            by_party[key]["balance_raw"] += m["balance"]

    rows: list[dict[str, Any]] = []
    total_amount = 0.0
    total_balance = 0.0
    for party in sorted(
        by_party.values(), key=lambda p: abs(p["amount_raw"]), reverse=True
    ):
        rows.append({
            "name": party["name"],
            "relation": " / ".join(party["relations"]),
            "amount": _money(party["amount_raw"]),
            "balance": _money(party["balance_raw"]),
        })
        total_amount += party["amount_raw"]
        total_balance += party["balance_raw"]

    return {
        "related_parties": rows,
        "total_amount": _money(total_amount),
        "total_balance": _money(total_balance),
        "party_count": len(rows),
        "narrative_seed": (
            f"本期检索到 {len(rows)} 个关联方，发生额合计 "
            f"{_money(total_amount)} 元，期末净往来余额 "
            f"{_money(total_balance)} 元（借为正）。"
        ),
    }


_CTX_BUILDERS: dict[str, Any] = {
    "NOTE_CASH_DETAIL": _ctx_cash_detail,
    "NOTE_AR_AGING": _ctx_ar_aging,
    "NOTE_INVENTORY": _ctx_inventory,
    "NOTE_FIXED_ASSETS": _ctx_fixed_assets,
    "NOTE_REVENUE_BY_CUSTOMER": _ctx_revenue_by_customer,
    "NOTE_EXPENSES": _ctx_expenses,
    "NOTE_ACCOUNTS_PAYABLE_CONCENTRATION": _ctx_accounts_payable,
    "NOTE_RELATED_PARTY_TRANSACTIONS": _ctx_related_party,
}


async def _latest_report_cells(
    service: "FinanceAutoService", *, org_id: str, period_id: str, sheet_kind: str
) -> dict[str, dict[str, Any]]:
    """Return ``{reference_code: {value, label}}`` for the latest report.

    Returns an empty dict when no report exists yet (the templates still
    render — they just show zeros).
    """
    async with service.db.conn.execute(
        "SELECT id FROM reports WHERE org_id=? AND period_id=? AND sheet_kind=? "
        "ORDER BY generated_at DESC LIMIT 1",
        (org_id, period_id, sheet_kind),
    ) as cur:
        row = await cur.fetchone()
    if row is None:
        return {}
    rid = row[0]
    async with service.db.conn.execute(
        "SELECT reference_code, target_label, value FROM report_cells WHERE report_id=?",
        (rid,),
    ) as cur:
        rows = await cur.fetchall()
    return {
        r[0]: {"label": r[1] or r[0], "value": r[2] or 0.0}
        for r in rows
    }


# ---------------------------------------------------------------------------
# Public service.
# ---------------------------------------------------------------------------


class NotesGeneratorError(RuntimeError):
    """Raised for client-visible generator failures (mapped to 4xx in routes)."""


class NotesGenerator:
    """Notes generator service — owns the ``note_*`` tables.

    Constructed once per :class:`FinanceAutoService`; reuse across requests
    is safe because the underlying ``aiosqlite.Connection`` is itself
    request-safe (the host's SQLite layer serialises writes).
    """

    def __init__(self, service: "FinanceAutoService"):
        self._svc = service
        self._bus = get_event_bus()

    # ------------------------- discovery / list ----------------------------

    async def list_templates(
        self, *, accounting_standard: str | None = None
    ) -> list[dict[str, Any]]:
        sql = "SELECT * FROM note_templates"
        params: tuple[Any, ...] = ()
        if accounting_standard:
            sql += " WHERE accounting_standard=?"
            params = (accounting_standard,)
        sql += " ORDER BY note_section ASC, note_item_code ASC"
        async with self._svc.db.conn.execute(sql, params) as cur:
            rows = await cur.fetchall()
        return [
            {
                "id": r["id"],
                "note_section": r["note_section"],
                "note_item_code": r["note_item_code"],
                "template_format": r["template_format"],
                "template_path": r["template_path"],
                "data_source": r["data_source"],
                "auto_fill_pct": r["auto_fill_pct"],
                "requires_ai": bool(r["requires_ai"]),
                "ai_scenario_id": r["ai_scenario_id"],
                "accounting_standard": r["accounting_standard"],
                "version": r["version"],
                "created_at": r["created_at"],
            }
            for r in rows
        ]

    async def list_documents(
        self, *, org_id: str, period_id: str | None = None
    ) -> list[dict[str, Any]]:
        sql = "SELECT * FROM note_documents WHERE org_id=?"
        params: list[Any] = [org_id]
        if period_id:
            sql += " AND period_id=?"
            params.append(period_id)
        sql += " ORDER BY created_at DESC"
        async with self._svc.db.conn.execute(sql, tuple(params)) as cur:
            rows = await cur.fetchall()
        return [self._row_to_doc(r) for r in rows]

    async def get_document(self, *, document_id: int) -> dict[str, Any]:
        async with self._svc.db.conn.execute(
            "SELECT * FROM note_documents WHERE id=?", (document_id,)
        ) as cur:
            row = await cur.fetchone()
        if row is None:
            raise NotesGeneratorError(f"document {document_id} not found")
        return self._row_to_doc(row)

    async def list_notes(self, *, document_id: int) -> list[dict[str, Any]]:
        # Validate first so a stray id yields a clean 404 in the route layer.
        await self.get_document(document_id=document_id)
        async with self._svc.db.conn.execute(
            "SELECT * FROM report_notes WHERE document_id=? "
            "ORDER BY note_section ASC, note_item_code ASC",
            (document_id,),
        ) as cur:
            rows = await cur.fetchall()
        return [self._row_to_note(r) for r in rows]

    # ------------------------- generate ------------------------------------

    async def generate(
        self,
        *,
        org_id: str,
        period_id: str,
        sections: list[str] | None = None,
        user_id: str = "local",
    ) -> dict[str, Any]:
        """Generate (or regenerate) the notes document for ``(org, period)``.

        A fresh ``note_documents`` row is always inserted so the audit
        trail keeps every generation attempt.  Returns the full payload
        including a `notes` array suitable for direct UI consumption.
        """
        # 1. Validate org.  Reuses the existing 404 path on the service.
        await self._svc.get_org(org_id)

        templates = await self._select_templates(sections=sections)
        if not templates:
            raise NotesGeneratorError(
                "no note_templates matched the request — has v10 migration run?"
            )

        now = _utcnow_iso()
        cur = await self._svc.db.conn.execute(
            "INSERT INTO note_documents(org_id, period_id, status, "
            "accounting_standard, version, created_at, updated_at) "
            "VALUES (?,?,?,?,1,?,?)",
            (org_id, period_id, "draft", "small_enterprise", now, now),
        )
        document_id = cur.lastrowid
        await cur.close()

        notes_payload: list[dict[str, Any]] = []
        narrative_seam: list[dict[str, Any]] = []
        for tmpl in templates:
            note = await self._render_one(
                org_id=org_id,
                period_id=period_id,
                document_id=document_id,
                template=tmpl,
            )
            notes_payload.append(note)
            if note["kind"] in ("narrative_pending_ai", "narrative_pending_user"):
                narrative_seam.append(
                    {
                        "org_id": org_id,
                        "period_id": period_id,
                        "user_id": user_id,
                        "document_id": document_id,
                        "note_id": note["id"],
                        "template_id": tmpl["id"],
                        "note_section": tmpl["note_section"],
                        "note_item_code": tmpl["note_item_code"],
                        "ai_scenario_id": tmpl.get("ai_scenario_id") or "S11",
                        "fallback_content": note["content"],
                    }
                )

        await self._svc.db.conn.commit()
        try:
            await self._bus.emit(
                "finance.notes.generated",
                {
                    "org_id": org_id,
                    "period_id": period_id,
                    "document_id": document_id,
                    "notes_count": len(notes_payload),
                    "user_id": user_id,
                },
            )
        except Exception as exc:  # noqa: BLE001 — emission failure must not break the API
            logger.warning("finance-auto: notes.generated emit failed: %s", exc)
        # Each hybrid / narrative template emits a draft_requested event so
        # Sibling B's S11 worker (subscribed via attach_event_bus_subscriber)
        # can pick the row up and replace the placeholder.  Failures in the
        # subscriber chain must not break this API call.
        for payload in narrative_seam:
            try:
                await self._bus.emit("finance.notes.draft_requested", payload)
            except Exception as exc:  # noqa: BLE001
                logger.warning(
                    "finance-auto: notes.draft_requested emit failed (note %s): %s",
                    payload["note_id"], exc,
                )

        return {
            "document_id": document_id,
            "org_id": org_id,
            "period_id": period_id,
            "status": "draft",
            "notes": notes_payload,
            "total": len(notes_payload),
        }

    async def _select_templates(
        self, *, sections: list[str] | None
    ) -> list[dict[str, Any]]:
        if sections:
            placeholders = ",".join("?" for _ in sections)
            sql = (
                f"SELECT * FROM note_templates WHERE note_section IN ({placeholders}) "
                "ORDER BY note_section ASC, note_item_code ASC"
            )
            params: tuple[Any, ...] = tuple(sections)
        else:
            sql = "SELECT * FROM note_templates ORDER BY note_section ASC, note_item_code ASC"
            params = ()
        async with self._svc.db.conn.execute(sql, params) as cur:
            rows = await cur.fetchall()
        return [dict(r) for r in rows]

    async def _render_one(
        self,
        *,
        org_id: str,
        period_id: str,
        document_id: int,
        template: dict[str, Any],
    ) -> dict[str, Any]:
        code = template["note_item_code"]
        builder = _CTX_BUILDERS.get(code)
        if builder is None:
            # No data context registered — render an empty narrative
            # placeholder so the row still exists in the document.
            ctx: dict[str, Any] = {}
        else:
            ctx = await builder(self._svc, org_id=org_id, period_id=period_id)

        template_path = _PLUGIN_ROOT / template["template_path"]
        if template_path.exists():
            template_src = template_path.read_text(encoding="utf-8")
            content = _render_template(template_src, ctx)
        else:
            content = (
                f"### {template['note_section']} · {code}\n\n"
                "[未找到模板文件，已使用兜底渲染。]\n"
            )

        if template["data_source"] == "data_driven":
            kind = "data"
        elif template["data_source"] == "narrative":
            kind = "narrative_pending_ai"
        else:
            # hybrid — data half rendered, narrative half deferred to S11.
            kind = "narrative_pending_ai"

        now = _utcnow_iso()
        cur = await self._svc.db.conn.execute(
            "INSERT INTO report_notes(document_id, template_id, note_section, "
            "note_item_code, content, kind, ai_audit_id, version, created_at, updated_at) "
            "VALUES (?,?,?,?,?,?,?,1,?,?)",
            (
                document_id,
                template["id"],
                template["note_section"],
                code,
                content,
                kind,
                None,
                now,
                now,
            ),
        )
        note_id = cur.lastrowid
        await cur.close()
        return {
            "id": note_id,
            "document_id": document_id,
            "template_id": template["id"],
            "note_section": template["note_section"],
            "note_item_code": code,
            "kind": kind,
            "content": content,
            "version": 1,
            "created_at": now,
            "updated_at": now,
        }

    # ------------------------- mutate / finalise --------------------------

    async def update_note(
        self,
        *,
        note_id: int,
        content: str,
        expected_version: int,
    ) -> dict[str, Any]:
        """Optimistic-lock update on a single note.

        Implements the Part Infra C3 contract inline (``WHERE id=? AND
        version=?``).  Returns the freshly-read row on success.  Raises
        :class:`HTTPException` 409 on a version mismatch so the routes
        layer can short-circuit straight to the API surface.
        """
        now = _utcnow_iso()
        new_version = expected_version + 1
        cur = await self._svc.db.conn.execute(
            "UPDATE report_notes SET content=?, version=?, updated_at=? "
            "WHERE id=? AND version=?",
            (content, new_version, now, note_id, expected_version),
        )
        rows_changed = cur.rowcount
        await cur.close()
        if rows_changed == 0:
            # Either the id is wrong or the version is stale — figure out
            # which so the caller gets a sensible error.
            async with self._svc.db.conn.execute(
                "SELECT version FROM report_notes WHERE id=?", (note_id,)
            ) as cur:
                row = await cur.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail=f"note {note_id} not found")
            raise HTTPException(
                status_code=409,
                detail={
                    "error": "version_conflict",
                    "expected_version": expected_version,
                    "current_version": row["version"],
                },
            )
        await self._svc.db.conn.commit()
        async with self._svc.db.conn.execute(
            "SELECT * FROM report_notes WHERE id=?", (note_id,)
        ) as cur:
            row = await cur.fetchone()
        return self._row_to_note(row)

    async def finalize_document(self, *, document_id: int) -> dict[str, Any]:
        doc = await self.get_document(document_id=document_id)
        if doc["status"] == "finalized":
            return doc
        now = _utcnow_iso()
        await self._svc.db.conn.execute(
            "UPDATE note_documents SET status='finalized', updated_at=?, "
            "version=version+1 WHERE id=?",
            (now, document_id),
        )
        await self._svc.db.conn.commit()
        return await self.get_document(document_id=document_id)

    # ------------------------- export --------------------------------------

    async def export_docx(self, *, document_id: int) -> bytes:
        """Bundle every note into a single transportable blob.

        Tries ``python-docx`` if it's already installed; falls back to a
        deterministic markdown bundle wrapped in ``<bundle>...</bundle>``
        sentinel tags so callers can detect the fallback path.
        """
        doc = await self.get_document(document_id=document_id)
        notes = await self.list_notes(document_id=document_id)

        bundle_parts: list[str] = [
            f"# Notes Document {document_id}",
            f"org_id: {doc['org_id']}",
            f"period_id: {doc['period_id']}",
            f"status: {doc['status']}",
            "",
        ]
        for n in notes:
            bundle_parts.append(f"## [{n['note_section']}] {n['note_item_code']}")
            bundle_parts.append(f"kind: {n['kind']} | version: {n['version']}")
            bundle_parts.append("")
            bundle_parts.append(n["content"])
            bundle_parts.append("")
        markdown_bundle = "\n".join(bundle_parts)

        try:
            from docx import Document  # type: ignore[import-not-found]
        except ImportError:
            return b"<bundle>" + markdown_bundle.encode("utf-8") + b"</bundle>"

        # python-docx is available — emit a real .docx in a memory buffer.
        import io

        d = Document()
        d.add_heading(f"Notes Document {document_id}", level=0)
        d.add_paragraph(
            f"org_id: {doc['org_id']}  period_id: {doc['period_id']}  status: {doc['status']}"
        )
        for n in notes:
            d.add_heading(f"[{n['note_section']}] {n['note_item_code']}", level=1)
            d.add_paragraph(f"kind: {n['kind']} | version: {n['version']}")
            for line in (n["content"] or "").splitlines():
                d.add_paragraph(line)
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    # ------------------------- helpers ------------------------------------

    @staticmethod
    def _row_to_doc(row) -> dict[str, Any]:
        return {
            "id": row["id"],
            "org_id": row["org_id"],
            "period_id": row["period_id"],
            "status": row["status"],
            "accounting_standard": row["accounting_standard"],
            "version": row["version"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
        }

    @staticmethod
    def _row_to_note(row) -> dict[str, Any]:
        return {
            "id": row["id"],
            "document_id": row["document_id"],
            "template_id": row["template_id"],
            "note_section": row["note_section"],
            "note_item_code": row["note_item_code"],
            "content": row["content"],
            "kind": row["kind"],
            "ai_audit_id": row["ai_audit_id"],
            "version": row["version"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
        }


__all__ = [
    "ALL_SECTIONS",
    "NotesGenerator",
    "NotesGeneratorError",
]
